from docx import Document as createdoc
from docx.document import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import re
import os


class Expreport:
    paras = []
    todo_paras = []

    def __init__(self, paras):
        self.doc = createdoc("static/template.docx")
        self.todo_paras = paras
        self.check_paras()

    def fill_all(self):
        self.it = iter(self.todo_paras)
        self.fill_base("姓名", "肖逸程")
        self.fill_base("号", "191304324")
        # self.check_paras()
        self.fill_paras()
        # self.doc.save("me.docx")

    def save(self, str):
        self.doc.save(str)

    def check_paras(self):
        self.paras.clear()
        self.ita_paras(lambda paras, i ,p :self.paras.append(paras[i-1].text))

    def fill_todo(self, paras, i, p):
        p.text = "\t" + next(self.it)
        modr = p.runs[0].font
        modr.size = Pt(14)
        modr.name = "宋体"
        r = p.runs[0]._element.rPr.rFonts
        r.set(qn('w:eastAsia'), "宋体")

    def fill_paras(self):
        self.ita_paras(self.fill_todo)

    def fill_base(self, title, istr):
        paras = self.doc.paragraphs
        for p in paras:
            if re.search(title, p.text):
                for run in p.runs:
                    if (run.text == "todo"):
                        run.text = istr

    def ita_paras(self, operation):
        paras = self.doc.paragraphs
        for i, p in enumerate(paras):
            if p.text == "todo":
                operation(paras, i, p)


# doc = createdoc("learnword.docx") #type: Document
# para = doc.paragraphs[0]
# for run in para.runs:
#     print(run.text + "..")


# document = Document()

# document.add_heading('Document Title', 0)

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631' 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

# document.save('demo2.docx')
